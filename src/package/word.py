import os
import platform
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pathlib import Path


class Word:
    def __init__(self):
        self.document = Document()

    def add_heading(self, text: str = "", taille: int = 1):
        self.entete = self.document.add_heading(text, taille)
        font = self.entete.style.font
        font.name = 'Times New Roman'
        font.color.rgb = RGBColor(0, 0, 0)
        para = self.entete.paragraph_format
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_paragraph(self, text: str = ""):
        paragraph = self.document.add_paragraph(text)
        format = paragraph.paragraph_format
        format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def break_new_section(self):
        self.document.add_page_break()

    def head_table(self, heads: list):
        self.col = len(heads)
        # tableau
        self.table = self.document.add_table(rows=1, cols=self.col)
        # self.table.style = 'Medium Shading 1 Accent 1'
        self.table.style = 'Table Grid'

        # entete de du tableau
        self.head_cells = self.table.rows[0].cells
        for i in range(len(heads)):
            self.head_cells[i].text = str(heads[i])

    def add_collections(self, items):
        if isinstance(items, list):
            li = 0
            for i in range(len(items)):
                li += 1
                if isinstance(items[i], list):
                    self.row_cells = self.table.add_row().cells
                    for j in range(len(items[i])):
                        self.row_cells[j].text = str(items[i][j])
                else:
                    self.row_cells = self.table.add_row().cells
                    for item in items:
                        self.row_cells[0].text = str(li)
                        for k in range(len(items[i].in_liste())-1):
                            self.row_cells[k+1].text = str(items[i].in_liste()[k])
        else:
            print("première list")

    def save(self):
        self.document.save(self.path)

    def open(self):
        if platform.system() == 'Darwin':
            os.system(f"open {self.path}")
        elif platform.system() in ["Windows", "windows", "cygwin", "win32"]:
            os.system(f"start {self.path}")
        else:
            os.system(f"./{self.path}")

    @property
    def path(self):
        return os.path.join(Path.home(), "repartition.docx")
